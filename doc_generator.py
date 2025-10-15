"""
Automated Documentation Generator - Gemini Edition

Creates comprehensive Word documents from codebase analysis using Google's Gemini API.

This system orchestrates four main agents to transform raw code into professional
technical documentation:
1. GeminiDocAgent - Interacts with Google's Gemini API for content generation
2. ScreenshotAgent - Captures visual representations of code and applications
3. DocumentAssembler - Builds formatted Word documents with embedded media
4. DocumentationGenerator - Coordinates the entire 4-phase generation pipeline

Why this exists:
- Manual documentation is time-consuming and often incomplete
- AI-generated docs provide consistent structure and comprehensive coverage
- Screenshots make documentation more accessible and visually informative
- Word format allows easy editing and distribution across teams

Architecture notes:
- Gemini API chosen over other LLMs for its 1M token context window (crucial for
  understanding large codebases) and strong technical writing capabilities
- python-docx for output format (Word documents are universally editable, unlike PDFs)
- Selenium for screenshots (ensures consistent rendering across platforms)
- Environment-based configuration for flexibility across different projects
"""

import os
import json
import time
import platform
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass, field

import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.firefox import GeckoDriverManager

# Load environment variables from .env file
# Why: Keeps sensitive API keys out of source code and allows easy configuration
# without code changes. Critical for security and multi-project usage.
load_dotenv()

# Mermaid diagram support
try:
    import subprocess
    MERMAID_CLI_AVAILABLE = True
except ImportError:
    MERMAID_CLI_AVAILABLE = False


@dataclass
class DocumentSection:
    """Represents a single section of the generated documentation.

    Each section is a hierarchical unit (heading + content + media) that will be
    rendered in the final Word document. Sections can be nested via the level attribute.

    Why this exists:
    - Provides structured data for the documentation generation pipeline
    - Separates concerns: content generation vs. document assembly
    - Allows sections to be processed independently (parallel generation in future)

    Attributes:
        title (str): Section heading text (e.g., "Installation Guide")
        level (int): Heading hierarchy (1=main, 2=sub, 3=sub-sub, etc.)
        content (str): Markdown-formatted text content from Gemini
        images (List[Dict]): Screenshot metadata [{"description": str, "path": str}]
        code_blocks (List[str]): Extracted code examples to be formatted separately

    Example:
        section = DocumentSection(
            title="API Reference",
            level=1,
            content="The API provides...",
            images=[{"description": "API endpoints", "path": "/path/to/screenshot.png"}]
        )
    """
    title: str
    level: int  # 1 for main heading, 2 for subheading, etc.
    content: str = ""
    images: List[Dict[str, str]] = field(default_factory=list)
    code_blocks: List[str] = field(default_factory=list)
    mermaid_diagrams: List[Dict[str, str]] = field(default_factory=list)  # {"description": str, "code": str, "path": str}


@dataclass
class DocumentationPlan:
    """Overall documentation structure created by AI analysis.

    The plan is the blueprint for the entire documentation, created in Phase 2 by
    analyzing the codebase. It defines what sections will exist and their hierarchy.

    Why this exists:
    - Separates planning (what to document) from execution (writing content)
    - Allows validation/modification of structure before expensive content generation
    - Enables consistent documentation structure across different projects

    Attributes:
        title (str): Document title (e.g., "MyProject Technical Documentation")
        sections (List[DocumentSection]): Ordered list of sections to generate

    Note:
        Sections are initially created with empty content; content is filled in Phase 3.
    """
    title: str
    sections: List[DocumentSection]


class GeminiDocAgent:
    """Handles all Google Gemini API interactions for documentation generation.

    This agent is responsible for:
    1. Creating documentation structure (sections, hierarchy) via codebase analysis
    2. Generating detailed content for each section with context awareness
    3. Identifying which code files/areas need screenshots for visual documentation

    Why Gemini API specifically:
    - 1M token context window allows understanding entire codebases at once
    - Superior technical writing quality compared to other LLMs
    - Free tier sufficient for most documentation projects (15 req/min)
    - Native JSON mode for structured outputs (documentation plans)

    Rate Limits (Free Tier - as of 2025):
    - 15 requests per minute
    - 1,500 requests per day
    - 1M tokens/minute input
    - 32K tokens/minute output

    Performance considerations:
    - Context truncation to stay within token limits (100K chars ≈ 25K tokens)
    - Request delay prevents rate limit errors (2 seconds between calls)
    - Retry logic handles transient API failures

    Security considerations:
    - API key loaded from environment, never hardcoded
    - No sensitive data sent to API (codebase only, no credentials)
    - API responses are not cached (privacy)

    Attributes:
        model_name (str): Gemini model identifier (e.g., 'gemini-2.5-pro')
        temperature (float): AI creativity setting (0.0=deterministic, 1.0=creative)
        max_tokens (int): Maximum response length per request
        request_delay (int): Seconds to wait between API calls (rate limiting)
        model (GenerativeModel): Configured Gemini model instance

    Example:
        agent = GeminiDocAgent()
        plan = agent.create_documentation_plan(codebase_text, "MyProject")
        content = agent.generate_section_content(plan.sections[0], codebase_text)
    """

    def __init__(self):
        """Initialize the Gemini API client with configuration from environment.

        Raises:
            ValueError: If GEMINI_API_KEY is not set in environment variables
        """
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")

        # Configure the global Gemini SDK with API key
        genai.configure(api_key=api_key)

        # Load model configuration from environment with sensible defaults
        # Why configurable: Different projects need different tradeoffs between
        # quality (gemini-2.5-pro) and speed/cost (gemini-2.0-flash-exp)
        self.model_name = os.getenv('GEMINI_MODEL', 'gemini-2.5-pro')
        self.temperature = float(os.getenv('GEMINI_TEMPERATURE', '0.7'))
        self.max_tokens = int(os.getenv('GEMINI_MAX_OUTPUT_TOKENS', '8000'))

        # Rate limiting configuration with safer defaults
        # Default 5 seconds = 12 req/min (safely under 15 req/min free tier limit)
        self.request_delay = int(os.getenv('GEMINI_REQUEST_DELAY', '5'))

        # Specific delays for different request types (all default to base delay if not set)
        self.plan_request_delay = int(os.getenv('GEMINI_PLAN_REQUEST_DELAY', str(self.request_delay)))
        self.section_request_delay = int(os.getenv('GEMINI_SECTION_REQUEST_DELAY', str(self.request_delay)))
        self.screenshot_request_delay = int(os.getenv('GEMINI_SCREENSHOT_REQUEST_DELAY', str(self.request_delay)))
        self.diagram_request_delay = int(os.getenv('GEMINI_DIAGRAM_REQUEST_DELAY', str(self.request_delay)))

        # Request tracking for per-minute rate limiting
        self.request_timestamps = []  # List of timestamps for recent requests
        self.max_requests_per_minute = int(os.getenv('GEMINI_MAX_REQUESTS_PER_MINUTE', '15'))

        # Exponential backoff configuration
        self.max_retries = int(os.getenv('GEMINI_MAX_RETRIES', '3'))
        self.base_backoff_delay = int(os.getenv('GEMINI_BASE_BACKOFF_DELAY', '5'))

        # Initialize the model with generation parameters
        # Why these settings:
        # - temperature=0.7: Balanced creativity (too low=repetitive, too high=random)
        # - max_tokens=8000: Sufficient for detailed sections without truncation
        self.model = genai.GenerativeModel(
            model_name=self.model_name,
            generation_config={
                'temperature': self.temperature,
                'max_output_tokens': self.max_tokens,
            }
        )

        print(f"✓ Initialized Gemini model: {self.model_name}")
        print(f"✓ Rate limiting: {self.request_delay}s base delay, max {self.max_requests_per_minute} req/min")
    
    def _track_request(self):
        """Track request timestamp and enforce per-minute rate limit.

        Removes timestamps older than 60 seconds and checks if we're at the limit.
        If at limit, waits until oldest request expires.
        """
        current_time = time.time()

        # Remove timestamps older than 60 seconds
        self.request_timestamps = [ts for ts in self.request_timestamps if current_time - ts < 60]

        # Check if we're at the request limit
        if len(self.request_timestamps) >= self.max_requests_per_minute:
            # Calculate how long to wait until the oldest request expires
            oldest_timestamp = self.request_timestamps[0]
            wait_time = 60 - (current_time - oldest_timestamp) + 1  # +1 for safety margin
            if wait_time > 0:
                print(f"⏳ Rate limit approaching ({len(self.request_timestamps)}/{self.max_requests_per_minute} req/min). Waiting {wait_time:.1f}s...")
                time.sleep(wait_time)
                # Recalculate current time after waiting
                current_time = time.time()
                # Remove expired timestamps again
                self.request_timestamps = [ts for ts in self.request_timestamps if current_time - ts < 60]

        # Add current request timestamp
        self.request_timestamps.append(current_time)

    def _make_request(self, prompt: str, delay: bool = True, json_mode: bool = False, request_type: str = 'default') -> str:
        """Make a request to Gemini API with intelligent rate limiting and exponential backoff.

        Enhanced rate limiting features:
        - Per-minute request tracking to stay under 15 req/min limit
        - Type-specific delays for different request types
        - Exponential backoff on 429 rate limit errors
        - Configurable retry logic with increasing delays

        Why intelligent rate limiting:
        - Free tier: 15 req/min, 1,500 req/day, 1M tokens/min input, 32K tokens/min output
        - Default 5s delay = 12 req/min (safe margin under 15 req/min)
        - Different request types can have custom delays (e.g., slower for plans, faster for sections)
        - Request tracking prevents bursts from exceeding limits

        Why exponential backoff:
        - 429 errors indicate we hit rate limits despite tracking
        - Exponential backoff (5s, 10s, 20s) gives API time to recover
        - Prevents cascading failures and wasted retry attempts

        Args:
            prompt (str): The prompt to send to Gemini
            delay (bool): Whether to apply rate limit delay before request.
                         Set to False for first request in a session to avoid
                         unnecessary wait.
            json_mode (bool): Whether to force JSON response format.
                            Set to True when expecting structured JSON output.
            request_type (str): Type of request for custom delays.
                              Options: 'plan', 'section', 'screenshot', 'diagram', 'default'

        Returns:
            str: The generated text response from Gemini

        Raises:
            Exception: If API fails after all retries (re-raises last exception)

        Example:
            # Plan creation with custom delay
            response = agent._make_request(prompt, delay=False, json_mode=True, request_type='plan')

            # Section content with standard delay
            response = agent._make_request(prompt, request_type='section')
        """
        # Select appropriate delay based on request type
        delay_time = self.request_delay  # default
        if request_type == 'plan':
            delay_time = self.plan_request_delay
        elif request_type == 'section':
            delay_time = self.section_request_delay
        elif request_type == 'screenshot':
            delay_time = self.screenshot_request_delay
        elif request_type == 'diagram':
            delay_time = self.diagram_request_delay

        # Rate limiting: Wait before making request to respect API limits
        if delay:
            time.sleep(delay_time)

        # Track this request for per-minute rate limiting
        self._track_request()

        # Build generation configuration
        generation_config = {
            'temperature': self.temperature,
            'max_output_tokens': self.max_tokens,
        }

        # Force JSON mode for structured outputs (documentation plans)
        if json_mode:
            generation_config['response_mime_type'] = 'application/json'

        # Exponential backoff retry logic
        last_exception = None
        for retry_attempt in range(self.max_retries):
            try:
                response = self.model.generate_content(
                    prompt,
                    generation_config=generation_config
                )
                return response.text

            except Exception as e:
                last_exception = e
                error_str = str(e).lower()

                # Check if this is a rate limit error (429)
                is_rate_limit_error = '429' in error_str or 'quota' in error_str or 'rate limit' in error_str

                if retry_attempt < self.max_retries - 1:
                    # Calculate exponential backoff delay
                    if is_rate_limit_error:
                        # Longer backoff for rate limit errors
                        backoff_delay = self.base_backoff_delay * (2 ** retry_attempt)
                        print(f"⚠️  Rate limit error (429): {e}")
                        print(f"   Attempt {retry_attempt + 1}/{self.max_retries}. Backing off for {backoff_delay}s...")
                    else:
                        # Standard backoff for other errors
                        backoff_delay = self.base_backoff_delay
                        print(f"⚠️  Gemini API error: {e}")
                        print(f"   Attempt {retry_attempt + 1}/{self.max_retries}. Retrying in {backoff_delay}s...")

                    time.sleep(backoff_delay)
                else:
                    # Final attempt failed
                    print(f"❌ All {self.max_retries} retry attempts failed")
                    raise last_exception

        # Should not reach here, but just in case
        raise last_exception
    
    def create_documentation_plan(self, context: str, project_name: str,
                                  min_sections: int = 9, max_sections: int = 15) -> DocumentationPlan:
        """Analyze codebase and create structured documentation outline (Phase 2).

        This is the first AI interaction in the pipeline. Gemini analyzes the entire
        codebase context and creates a comprehensive documentation structure with
        sections, hierarchy, and placeholders for images.

        Why this approach:
        - Separating planning from content generation allows better organization
        - AI can see the big picture before diving into details
        - Structured JSON output enables programmatic processing
        - User can validate/modify structure before expensive content generation

        Why JSON output:
        - Structured data is easier to parse than freeform text
        - Enables validation (does plan have required sections?)
        - Allows programmatic manipulation (reorder, filter sections)

        Args:
            context (str): Full codebase text (from repomix or directory scan)
            project_name (str): Human-readable project name for documentation
            min_sections (int): Minimum number of sections to generate
            max_sections (int): Maximum number of sections to generate

        Returns:
            DocumentationPlan: Structured plan with title and sections

        Note:
            - Context is truncated to 100K chars to stay within token limits
            - Gemini sometimes wraps JSON in markdown code blocks despite instructions
            - Fallback plan ensures generation continues even if AI output is malformed
        """

        # Context truncation for token limits
        # Why 100,000 chars: Rough estimate 1 token ≈ 4 chars, so 100K ≈ 25K tokens
        # This leaves room for prompt text and output while staying under 1M token limit
        # Performance consideration: Larger context = better AI understanding but
        # slower processing and higher risk of hitting token limits
        context = context[:100000]

        # Determine section count based on project size
        context_size = len(context)
        if context_size < 30000:
            # Small project: fewer sections
            actual_min_sections = max(5, min_sections - 4)
            actual_max_sections = max(9, max_sections - 6)
            section_guidance = f"This is a small project. Create {actual_min_sections}-{actual_max_sections} sections total"
        else:
            # Normal/large project: standard range
            actual_min_sections = min_sections
            actual_max_sections = max_sections
            section_guidance = f"Create {actual_min_sections}-{actual_max_sections} sections total"
        
        prompt = f"""You are a technical documentation expert. Analyze this codebase and create a comprehensive documentation plan.

Project: {project_name}

Codebase Context:
{context}

Create a detailed outline for technical documentation. You MUST return a JSON object with EXACTLY this structure at the root level - do NOT wrap it in any outer object:

{{
    "title": "Project Documentation Title",
    "sections": [
        {{
            "title": "Section Title",
            "level": 1,
            "needs_images": false,
            "image_descriptions": []
        }}
    ]
}}

CRITICAL REQUIREMENTS:
1. Return ONLY the JSON object shown above - no wrapper objects, no extra keys at root level
2. The root object MUST have exactly two keys: "title" and "sections"
3. Do NOT wrap this in "documentation_plan", "plan", or any other outer key
4. Each section MUST have: "title" (string), "level" (1 or 2), "needs_images" (boolean), "image_descriptions" (array)

{section_guidance}. Choose the appropriate sections based on project complexity:

Core sections (always include):
1. Overview/Introduction (level 1)
2. Installation & Setup (level 1)
3. Usage Guide (level 1) - with subsections if needed

Additional sections (include based on project needs and size):
4. Architecture & Design (level 1) - with 2-3 level 2 subsections for larger projects
5. Core Components (level 1) - with 2-4 level 2 subsections for major components
6. Configuration (level 1)
7. API Reference (level 1) - if applicable
8. Development Guide (level 1)
9. Troubleshooting (level 1)

For small projects, focus on essential sections. For larger projects, include comprehensive coverage.

Return ONLY the JSON object, nothing else."""

        # Make API request with JSON mode enabled
        # json_mode=True forces Gemini to return valid JSON instead of prose
        # delay=False because this is typically the first request in the pipeline
        # request_type='plan' uses plan-specific delay configuration
        response = self._make_request(prompt, delay=False, json_mode=True, request_type='plan')

        # JSON cleanup: Gemini API quirk
        # Problem: Despite json_mode and "return ONLY JSON" in prompt, Gemini occasionally
        # still wraps responses in markdown code blocks: ```json\n{...}\n```
        # Why this happens: Model trained on markdown-formatted conversations
        # Solution: Robustly strip code block markers before parsing
        response = response.strip()

        # Remove markdown code block markers if present
        if response.startswith('```'):
            # Find the first line break (end of opening marker)
            first_newline = response.find('\n')
            if first_newline != -1:
                response = response[first_newline + 1:]  # Skip first line (```json or ```)

            # Remove closing ``` if present at the end
            if response.rstrip().endswith('```'):
                # Find last occurrence of ```
                last_backticks = response.rstrip().rfind('```')
                response = response[:last_backticks]

        response = response.strip()

        # Parse JSON with fallback for robustness
        try:
            plan_json = json.loads(response)

            # Handle nested response: Gemini sometimes wraps the response in an outer key
            # Check for common wrapper keys and unwrap if found
            if len(plan_json) == 1 and "documentation_plan" in plan_json:
                print("   Unwrapping nested 'documentation_plan' key")
                plan_json = plan_json["documentation_plan"]
            elif len(plan_json) == 1 and "plan" in plan_json:
                print("   Unwrapping nested 'plan' key")
                plan_json = plan_json["plan"]

            # Normalize key names: Gemini sometimes uses alternative key names
            # Handle "planTitle" or "title" for document title
            if "planTitle" in plan_json and "title" not in plan_json:
                plan_json["title"] = plan_json["planTitle"]

            # Handle "project_name" as title fallback
            if "title" not in plan_json and "project_name" in plan_json:
                plan_json["title"] = f"{plan_json['project_name']} Documentation"

            # Handle "documentationSections" or "sections" for sections array
            if "documentationSections" in plan_json and "sections" not in plan_json:
                plan_json["sections"] = plan_json["documentationSections"]

            # Handle other common variations
            if "chapters" in plan_json and "sections" not in plan_json:
                plan_json["sections"] = plan_json["chapters"]
            if "content_sections" in plan_json and "sections" not in plan_json:
                plan_json["sections"] = plan_json["content_sections"]
            if "documentation_sections" in plan_json and "sections" not in plan_json:
                plan_json["sections"] = plan_json["documentation_sections"]

            # Validate required keys exist after normalization
            # If either key is missing, treat as invalid response
            if "title" not in plan_json or "sections" not in plan_json:
                print(f"⚠️  Invalid plan structure: missing required keys")
                print(f"   Available keys: {list(plan_json.keys())}")
                # Print full JSON for debugging (pretty printed)
                print(f"   Full response structure:")
                print(json.dumps(plan_json, indent=2)[:1000])
                # Trigger fallback by raising KeyError
                raise KeyError("Missing required keys in plan JSON")

        except (json.JSONDecodeError, KeyError) as e:
            # Error handling: If AI returns invalid JSON, use basic fallback structure
            # Why fallback instead of failing: Better to have basic documentation than none
            # This can happen if: AI misunderstands prompt, network corruption, or API issues
            print(f"⚠️  JSON parsing error: {e}")
            print(f"Response was: {response[:500]}")

            # Fallback to basic documentation structure
            # These three sections are universally applicable to any software project
            plan_json = {
                "title": f"{project_name} Documentation",
                "sections": [
                    {"title": "Overview", "level": 1, "needs_images": False, "image_descriptions": []},
                    {"title": "Installation", "level": 1, "needs_images": False, "image_descriptions": []},
                    {"title": "Usage", "level": 1, "needs_images": True, "image_descriptions": ["example usage"]},
                ]
            }

        return DocumentationPlan(
            title=plan_json["title"],
            sections=[DocumentSection(
                title=s["title"],
                level=s["level"],
                content="",
                images=[{"description": desc, "path": ""} for desc in s.get("image_descriptions", [])],
                code_blocks=[]
            ) for s in plan_json["sections"]]
        )
    
    def generate_section_content(self, section: DocumentSection, context: str, 
                                 previous_sections: str = "") -> str:
        """Phase 2: Generate detailed content for each section"""
        
        # Limit context
        context = context[:80000]
        previous_sections = previous_sections[-10000:]
        
        prompt = f"""Generate concise, professional documentation content for this section targeted at business stakeholders and technical managers.

Section Title: {section.title}
Section Level: {section.level}

Full Project Context:
{context}

Previously Written Sections (for continuity):
{previous_sections}

CRITICAL REQUIREMENTS:
1. Write CONCISE content (200-400 words for level 1 sections, 150-250 words for level 2 subsections)
2. Focus on WHAT the system does and WHY, not implementation details
3. DO NOT include code snippets unless absolutely essential for setup/configuration
4. For technical concepts, describe them conceptually - screenshots will show implementation
5. Use clear, professional language suitable for stakeholders
6. If this section needs images, reference them with [IMAGE: description] placeholders
7. Use markdown formatting (**, *, bullet points) for readability
8. Focus on benefits, capabilities, and high-level architecture

AVOID:
- Long code examples (use screenshots instead)
- Implementation details (focus on concepts)
- Repetitive information
- Unnecessary technical jargon

Generate ONLY the section content. No preamble, no explanations about the content."""

        content = self._make_request(prompt, request_type='section')
        
        # Extract code blocks
        section.code_blocks = self._extract_code_blocks(content)
        
        return content
    
    def _extract_code_blocks(self, content: str) -> List[str]:
        """Extract code blocks from content"""
        blocks = []
        lines = content.split('\n')
        in_block = False
        current_block = []
        
        for line in lines:
            if 'CODE_BLOCK_START' in line or line.strip().startswith('```'):
                in_block = True
                # Skip language identifier in ```python style
                if line.strip().startswith('```') and len(line.strip()) > 3:
                    continue
                continue
            elif 'CODE_BLOCK_END' in line or (in_block and line.strip() == '```'):
                if current_block:
                    blocks.append('\n'.join(current_block))
                current_block = []
                in_block = False
                continue
            
            if in_block:
                current_block.append(line)
        
        return blocks
    
    def identify_screenshot_targets(self, section: DocumentSection, context: str) -> List[Dict]:
        """Identify what files/URLs to screenshot for a section"""
        
        if not section.images:
            return []
        
        context = context[:30000]
        
        prompt = f"""For this documentation section, identify specific screenshot targets.

Section: {section.title}
Images Needed: {[img['description'] for img in section.images]}

Project Context:
{context}

For each image, provide:
1. target_type: "code_file", "directory_structure", or "config_file"
2. target_path: specific file path relative to project root
3. instructions: brief note on what to capture

Return ONLY a JSON array (no markdown, no code blocks):
[
    {{
        "description": "image description",
        "target_type": "code_file",
        "target_path": "src/main.py",
        "instructions": "Focus on the main function"
    }}
]

Return only valid JSON."""

        response = self._make_request(prompt, json_mode=True, request_type='screenshot')

        # Clean response (remove markdown code blocks if present)
        response = response.strip()
        if response.startswith('```'):
            first_newline = response.find('\n')
            if first_newline != -1:
                response = response[first_newline + 1:]
            if response.rstrip().endswith('```'):
                last_backticks = response.rstrip().rfind('```')
                response = response[:last_backticks]
        response = response.strip()

        try:
            return json.loads(response)
        except json.JSONDecodeError as e:
            print(f"⚠️  Failed to parse screenshot targets JSON: {e}")
            print(f"Response was: {response[:200]}")
            return []


class ScreenshotAgent:
    """Handles automated screenshot capture"""
    
    def __init__(self):
        self.project_path = Path(os.getenv('PROJECT_PATH', '.'))
        self.screenshot_dir = Path(os.getenv('SCREENSHOTS_DIRECTORY', './screenshots'))
        self.screenshot_dir.mkdir(exist_ok=True, parents=True)
        self.browser = os.getenv('BROWSER_CHOICE', 'chrome').lower()
        self.wait_time = int(os.getenv('SCREENSHOT_WAIT_TIME', '3'))
        
        print(f"✓ Screenshot directory: {self.screenshot_dir.absolute()}")
    
    def _get_driver(self):
        """Get appropriate browser driver"""
        if self.browser == "chrome":
            options = ChromeOptions()
            options.add_argument('--headless=new')
            options.add_argument('--window-size=1200,800')
            options.add_argument('--disable-gpu')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            service = ChromeService(ChromeDriverManager().install())
            return webdriver.Chrome(service=service, options=options)
        
        elif self.browser == "firefox":
            options = FirefoxOptions()
            options.add_argument('--headless')
            options.add_argument('--width=1200')
            options.add_argument('--height=800')
            service = FirefoxService(GeckoDriverManager().install())
            return webdriver.Firefox(service=service, options=options)
        
        else:
            raise ValueError(f"Unsupported browser: {self.browser}")
    
    def capture_code_file(self, file_path: str, instructions: str = "") -> Optional[str]:
        """Capture screenshot of a code file"""
        full_path = self.project_path / file_path
        
        if not full_path.exists():
            print(f"⚠️  File not found: {full_path}")
            return None
        
        # Read file
        try:
            with open(full_path, 'r', encoding='utf-8') as f:
                code = f.read()
        except Exception as e:
            print(f"⚠️  Could not read {file_path}: {e}")
            return None
        
        # Limit code size
        max_lines = int(os.getenv('MAX_CODE_BLOCK_LINES', '50'))
        lines = code.split('\n')
        if len(lines) > max_lines:
            code = '\n'.join(lines[:max_lines]) + f"\n\n... (truncated, {len(lines) - max_lines} more lines)"
        
        # Create HTML with syntax highlighting
        html_content = f"""
        <html>
        <head>
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github-dark.min.css">
            <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/highlight.min.js"></script>
            <style>
                body {{ 
                    margin: 20px; 
                    background: #0d1117; 
                    font-family: 'Consolas', 'Monaco', monospace;
                }}
                pre {{ 
                    margin: 0; 
                    border-radius: 6px;
                    padding: 16px;
                }}
                code {{ 
                    font-family: 'Consolas', 'Monaco', monospace; 
                    font-size: 13px;
                    line-height: 1.5;
                }}
                .file-header {{
                    color: #c9d1d9;
                    font-size: 14px;
                    margin-bottom: 10px;
                    padding: 8px;
                    background: #161b22;
                    border-radius: 6px;
                }}
            </style>
        </head>
        <body>
            <div class="file-header">📄 {file_path}</div>
            <pre><code class="language-python">{code}</code></pre>
            <script>hljs.highlightAll();</script>
        </body>
        </html>
        """
        
        # Save temp HTML
        temp_html = self.screenshot_dir / "temp_code.html"
        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Capture screenshot
        try:
            driver = self._get_driver()
            driver.get(f"file://{temp_html.absolute()}")
            time.sleep(self.wait_time)
            
            screenshot_path = self.screenshot_dir / f"{file_path.replace('/', '_').replace('\\', '_')}.png"
            driver.save_screenshot(str(screenshot_path))
            driver.quit()
            
            return str(screenshot_path)
        except Exception as e:
            print(f"⚠️  Screenshot failed for {file_path}: {e}")
            return None
    
    def capture_directory_tree(self, base_path: str = ".") -> Optional[str]:
        """Create a visual directory structure"""
        full_path = self.project_path / base_path
        
        excluded = os.getenv('EXCLUDED_DIRECTORIES', '').split(',')
        excluded = [d.strip() for d in excluded if d.strip()]
        
        def build_tree(path: Path, prefix: str = "", is_last: bool = True, depth: int = 0) -> str:
            if depth > 4:  # Limit depth
                return ""
            
            if path.is_file():
                return f"{prefix}{'└── ' if is_last else '├── '}{path.name}\n"
            
            tree = f"{prefix}{'└── ' if is_last else '├── '}{path.name}/\n"
            
            try:
                children = sorted(path.iterdir(), key=lambda x: (x.is_file(), x.name))
                children = [c for c in children if c.name not in excluded][:20]
            except PermissionError:
                return tree
            
            for i, child in enumerate(children):
                is_last_child = i == len(children) - 1
                extension = "    " if is_last else "│   "
                tree += build_tree(child, prefix + extension, is_last_child, depth + 1)
            
            return tree
        
        tree_text = build_tree(full_path)
        
        # Create HTML
        html = f"""
        <html>
        <head>
            <style>
                body {{ 
                    background: #0d1117; 
                    color: #c9d1d9;
                    font-family: 'Consolas', 'Monaco', monospace;
                    padding: 20px;
                    font-size: 14px;
                    line-height: 1.6;
                }}
                pre {{ 
                    margin: 0;
                    background: #161b22;
                    padding: 16px;
                    border-radius: 6px;
                }}
                .header {{
                    color: #58a6ff;
                    font-size: 16px;
                    margin-bottom: 12px;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            <div class="header">📁 Project Structure</div>
            <pre>{tree_text}</pre>
        </body>
        </html>
        """
        
        temp_html = self.screenshot_dir / "temp_tree.html"
        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html)
        
        try:
            driver = self._get_driver()
            driver.get(f"file://{temp_html.absolute()}")
            time.sleep(self.wait_time)
            
            screenshot_path = self.screenshot_dir / "directory_structure.png"
            driver.save_screenshot(str(screenshot_path))
            driver.quit()
            
            return str(screenshot_path)
        except Exception as e:
            print(f"⚠️  Directory tree screenshot failed: {e}")
            return None
    
    def capture_live_url(self, url: str, name: str) -> Optional[str]:
        """Capture screenshot of a running application"""
        try:
            driver = self._get_driver()
            driver.get(url)
            time.sleep(self.wait_time)
            
            screenshot_path = self.screenshot_dir / f"live_{name}.png"
            driver.save_screenshot(str(screenshot_path))
            driver.quit()
            
            print(f"    ✓ Captured: {url}")
            return str(screenshot_path)
        except Exception as e:
            print(f"    ✗ Failed to capture {url}: {e}")
            return None


class MermaidAgent:
    """Handles Mermaid diagram generation for architecture visualization.

    This agent generates Mermaid diagrams to visualize:
    - System architecture (component diagrams)
    - Data flow (flowcharts, sequence diagrams)
    - Class hierarchies (class diagrams)
    - State machines (state diagrams)
    - Database schemas (ER diagrams)

    Why Mermaid:
    - Text-based diagram format (version control friendly)
    - Wide support (GitHub, GitLab, documentation tools)
    - Multiple diagram types (flowchart, sequence, class, state, ER, etc.)
    - Can be rendered to PNG/SVG for Word documents

    Rendering options:
    1. mermaid-cli (mmdc) - Node.js based, best quality
    2. Online API (mermaid.ink) - No installation required
    3. Fallback: Include Mermaid code as formatted text
    """

    def __init__(self):
        """Initialize Mermaid diagram generator."""
        self.diagrams_dir = Path(os.getenv('MERMAID_DIAGRAMS_DIRECTORY', './mermaid_diagrams'))
        self.diagrams_dir.mkdir(exist_ok=True, parents=True)
        self.use_mermaid = os.getenv('ENABLE_MERMAID_DIAGRAMS', 'true').lower() == 'true'

        # Check if mermaid-cli is installed
        self.mmdc_available = self._check_mmdc_available()

        print(f"✓ Mermaid diagrams directory: {self.diagrams_dir.absolute()}")
        if self.mmdc_available:
            print("✓ mermaid-cli (mmdc) is available")
        else:
            print("⚠️  mermaid-cli not found, will use fallback rendering")

    def _check_mmdc_available(self) -> bool:
        """Check if mermaid-cli (mmdc command) is available."""
        try:
            result = subprocess.run(['mmdc', '--version'],
                                   capture_output=True,
                                   timeout=5)
            return result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False

    def generate_diagram_code(self, gemini_agent, context: str, diagram_type: str,
                              description: str) -> Optional[str]:
        """Generate Mermaid diagram code using Gemini AI.

        Args:
            gemini_agent: GeminiDocAgent instance for API calls
            context: Project context for understanding architecture
            diagram_type: Type of diagram (flowchart, sequence, class, etc.)
            description: What the diagram should show

        Returns:
            str: Mermaid diagram code, or None if generation fails
        """
        prompt = f"""Generate a Mermaid diagram for technical documentation.

Diagram Type: {diagram_type}
Description: {description}

Project Context:
{context[:50000]}

Generate ONLY the Mermaid diagram code (no markdown code blocks, no explanations).
Start directly with the diagram type (e.g., 'graph TD', 'sequenceDiagram', 'classDiagram').

For {diagram_type}:
- Make it clear and readable
- Include relevant components/functions from the codebase
- Use proper Mermaid syntax
- Keep it focused (5-15 nodes maximum)
- Use descriptive labels

Return ONLY the Mermaid code."""

        try:
            response = gemini_agent._make_request(prompt, request_type='diagram')

            # Clean up response
            response = response.strip()
            # Remove markdown code blocks if present
            if response.startswith('```mermaid'):
                lines = response.split('\n')
                response = '\n'.join(lines[1:-1] if lines[-1].strip() == '```' else lines[1:])
            response = response.replace('```mermaid', '').replace('```', '').strip()

            return response
        except Exception as e:
            print(f"⚠️  Failed to generate Mermaid diagram: {e}")
            return None

    def render_diagram(self, mermaid_code: str, output_name: str) -> Optional[str]:
        """Render Mermaid code to PNG image.

        Args:
            mermaid_code: Mermaid diagram code
            output_name: Output filename (without extension)

        Returns:
            str: Path to generated PNG file, or None if rendering fails
        """
        output_path = self.diagrams_dir / f"{output_name}.png"

        # Strategy 1: Use mermaid-cli (mmdc) if available
        if self.mmdc_available:
            try:
                # Save Mermaid code to temporary file
                temp_mmd = self.diagrams_dir / f"{output_name}.mmd"
                with open(temp_mmd, 'w', encoding='utf-8') as f:
                    f.write(mermaid_code)

                # Render using mmdc
                cmd = [
                    'mmdc',
                    '-i', str(temp_mmd),
                    '-o', str(output_path),
                    '-b', 'transparent',  # Transparent background
                    '-t', 'default',  # Theme
                    '-w', '1200',  # Width
                    '-H', '800'  # Height
                ]

                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                if result.returncode == 0 and output_path.exists():
                    print(f"    ✓ Rendered: {output_name}")
                    # Clean up temp file
                    temp_mmd.unlink()
                    return str(output_path)
                else:
                    print(f"    ⚠️  mmdc failed: {result.stderr[:200]}")
            except Exception as e:
                print(f"    ⚠️  Rendering error: {e}")

        # Strategy 2: Use mermaid.ink online service
        try:
            import base64
            import urllib.request
            import urllib.parse

            # Encode Mermaid code
            encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
            url = f"https://mermaid.ink/img/{encoded}"

            # Download rendered image
            urllib.request.urlretrieve(url, output_path)

            if output_path.exists() and output_path.stat().st_size > 0:
                print(f"    ✓ Rendered via mermaid.ink: {output_name}")
                return str(output_path)
        except Exception as e:
            print(f"    ⚠️  mermaid.ink rendering failed: {e}")

        # Strategy 3: Save Mermaid code as text file (fallback)
        try:
            text_path = self.diagrams_dir / f"{output_name}.txt"
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(f"Mermaid Diagram Code:\n\n{mermaid_code}")
            print(f"    ℹ️  Saved Mermaid code to: {text_path}")
            return None  # No image available
        except Exception as e:
            print(f"    ⚠️  Failed to save Mermaid code: {e}")
            return None


class DocumentAssembler:
    """Assembles the final Word document"""
    
    def __init__(self):
        self.doc = Document()
        output_dir = Path(os.getenv('OUTPUT_DIRECTORY', './output'))
        output_dir.mkdir(exist_ok=True, parents=True)
        
        filename = os.getenv('OUTPUT_FILENAME', 'documentation.docx')
        self.output_path = output_dir / filename
        
        self._setup_styles()
        print(f"✓ Output will be saved to: {self.output_path.absolute()}")
    
    def _setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def add_title_page(self, title: str, project_name: str):
        """Create a professional title page"""
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        
        self.doc.add_paragraph()
        
        subtitle_para = self.doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"Project: {project_name}")
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        desc = os.getenv('PROJECT_DESCRIPTION', '')
        if desc:
            desc_para = self.doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_para.add_run(desc)
            desc_run.font.size = Pt(12)
            desc_run.font.italic = True
        
        self.doc.add_paragraph()
        
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {time.strftime('%B %d, %Y')}")
        date_run.font.size = Pt(11)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add author/organization if specified in environment
        author = os.getenv('DOCUMENTATION_AUTHOR', '')
        if author:
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Author: {author}")
            author_run.font.size = Pt(10)
            author_run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_page_break()
    
    def add_section(self, section: DocumentSection):
        """Add a complete section to the document"""
        # Add heading
        heading = self.doc.add_heading(section.title, level=section.level)
        if section.level == 1:
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        # Process content
        paragraphs = section.content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Handle image placeholders
            if '[IMAGE:' in para_text:
                parts = para_text.split('[IMAGE:')
                
                # Text before image
                if parts[0].strip():
                    self._add_formatted_paragraph(parts[0])
                
                # Extract image description and find matching image
                for part in parts[1:]:
                    img_desc = part.split(']')[0].strip()
                    text_after = ']'.join(part.split(']')[1:]).strip()
                    
                    # Find matching screenshot
                    matching_img = None
                    for img in section.images:
                        if img['description'].lower() in img_desc.lower() or img_desc.lower() in img['description'].lower():
                            matching_img = img
                            break
                    
                    if matching_img and matching_img.get('path') and os.path.exists(matching_img['path']):
                        self.doc.add_paragraph(f"Figure: {img_desc}", style='Caption')
                        try:
                            self.doc.add_picture(matching_img['path'], width=Inches(6))
                        except:
                            self.doc.add_paragraph(f"[Image: {img_desc}]")
                        self.doc.add_paragraph()
                    
                    if text_after:
                        self._add_formatted_paragraph(text_after)
            else:
                self._add_formatted_paragraph(para_text)
        
        # Add code blocks
        for code in section.code_blocks:
            self._add_code_block(code)
        
        # Add any remaining images
        for img in section.images:
            if img.get('path') and os.path.exists(img['path']):
                # Check if already added above
                already_added = False
                for para in self.doc.paragraphs[-10:]:
                    if img['description'] in para.text:
                        already_added = True
                        break
                
                if not already_added:
                    self.doc.add_paragraph(f"Figure: {img['description']}", style='Caption')
                    try:
                        self.doc.add_picture(img['path'], width=Inches(6))
                    except:
                        pass
                    self.doc.add_paragraph()
    
    def _add_formatted_paragraph(self, text: str):
        """Add paragraph with markdown formatting"""
        para = self.doc.add_paragraph()
        
        # Simple markdown parsing
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Bold
                run = para.add_run(part)
                run.font.bold = True
            else:
                italic_parts = part.split('*')
                for j, ipart in enumerate(italic_parts):
                    run = para.add_run(ipart)
                    if j % 2 == 1:
                        run.font.italic = True
    
    def _add_code_block(self, code: str):
        """Add formatted code block"""
        para = self.doc.add_paragraph(style='Normal')
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)
    
    def save(self):
        """Save the document"""
        self.doc.save(self.output_path)
        print(f"✅ Documentation saved to: {self.output_path.absolute()}")
        return str(self.output_path.absolute())



    def save_as_pdf(self, docx_path: Optional[str] = None) -> Optional[str]:
        """Convert the Word document to PDF format with cross-platform support.

        This method attempts to convert the .docx file to .pdf using multiple strategies:

        1. **python-docx2pdf (Windows only)**: Best quality, native Word rendering
           - Uses Microsoft Word COM interface on Windows
           - Preserves all formatting, images, styles perfectly
           - Limitation: Only works on Windows with Word installed

        2. **LibreOffice (Cross-platform fallback)**: Good quality, widely available
           - Uses LibreOffice command-line interface (headless mode)
           - Works on Linux, macOS, Windows
           - Requires LibreOffice to be installed
           - Quality: ~95% as good as native Word (minor font differences)

        3. **reportlab (Last resort fallback)**: Basic quality, pure Python
           - No external dependencies (always available if installed)
           - Limited formatting support (no complex styles/images)
           - Only recommended when other methods unavailable
           - Note: Not implemented in this version (placeholder for future)

        **Why this architecture?**
        - python-docx2pdf provides best quality but Windows-only (COM limitation)
        - LibreOffice provides excellent cross-platform alternative
        - Graceful degradation: tries best option first, falls back as needed
        - User gets clear error messages if all methods fail

        **Platform Compatibility:**
        - Windows: python-docx2pdf (primary) → LibreOffice (fallback)
        - Linux: LibreOffice (primary) → manual conversion message
        - macOS: LibreOffice (primary) → manual conversion message

        **Error Handling:**
        - Missing library: Prints installation instructions, returns None
        - Conversion failure: Prints error details, returns None
        - Missing LibreOffice: Suggests installation command for user's OS
        - Always non-fatal: PDF generation failure won't crash main program

        Args:
            docx_path (Optional[str]): Path to .docx file to convert. If None,
                                      uses the document saved by save() method.
                                      Default: None (use self.output_path)

        Returns:
            Optional[str]: Absolute path to generated .pdf file if successful,
                          None if conversion failed or libraries unavailable

        Raises:
            None: All exceptions are caught and logged. Method returns None on failure.

        Example:
            # Basic usage
            assembler = DocumentAssembler()
            assembler.add_title_page("Docs", "Project")
            docx_path = assembler.save()
            pdf_path = assembler.save_as_pdf()
            if pdf_path:
                print(f"PDF created: {pdf_path}")

            # Convert existing file
            pdf_path = assembler.save_as_pdf("/path/to/existing.docx")

        Environment Variables:
            ENABLE_PDF_EXPORT (bool): Set to 'true' to enable PDF generation
                                     Default: false (opt-in feature)
            PDF_EXPORT_METHOD (str): Force specific method: 'docx2pdf', 'libreoffice'
                                    Default: auto-detect best available method

        Installation Instructions (printed on error):
            Windows: pip install python-docx2pdf
            Linux: sudo apt install libreoffice-writer  # or yum/pacman
            macOS: brew install --cask libreoffice
        """
        # Check if PDF export is enabled via environment variable
        # Default to false to avoid surprising users with unexpected PDF generation
        pdf_enabled = os.getenv('ENABLE_PDF_EXPORT', 'false').lower() == 'true'
        if not pdf_enabled:
            print("ℹ️  PDF export disabled. Set ENABLE_PDF_EXPORT=true in .env to enable.")
            return None

        # Determine source .docx file path
        if docx_path is None:
            docx_path = str(self.output_path.absolute())
        else:
            docx_path = str(Path(docx_path).absolute())

        # Verify source file exists
        if not os.path.exists(docx_path):
            print(f"❌ PDF Export Error: Source file not found: {docx_path}")
            print("   Hint: Call save() before save_as_pdf() to create the .docx file")
            return None

        # Generate PDF output path (same directory, .pdf extension)
        pdf_path = Path(docx_path).with_suffix('.pdf')

        print(f"\n📄 Converting to PDF: {Path(docx_path).name} → {pdf_path.name}")

        # Detect operating system for platform-specific guidance
        current_os = platform.system()  # Returns: 'Windows', 'Linux', 'Darwin' (macOS)

        # Strategy 1: Try python-docx2pdf (Windows-optimized, best quality)
        # This is the preferred method on Windows as it uses native Word COM interface
        try:
            from docx2pdf import convert

            print("   Using: python-docx2pdf (native Word conversion)")
            print("   Quality: Excellent (100% formatting preserved)")

            # Attempt conversion
            # Note: convert() uses Microsoft Word COM interface on Windows
            # It will raise an exception on non-Windows platforms
            convert(docx_path, str(pdf_path))

            # Verify PDF was created successfully
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                print(f"✅ PDF successfully created: {pdf_path.absolute()}")
                return str(pdf_path.absolute())
            else:
                print("⚠️  python-docx2pdf completed but PDF file is missing/empty")
                # Fall through to next strategy

        except ImportError:
            # Library not installed - provide helpful installation guidance
            print("   python-docx2pdf not installed (skipping)")
            if current_os == 'Windows':
                print("   💡 For best PDF quality on Windows:")
                print("      pip install python-docx2pdf")
                print("      Requires: Microsoft Word installed")
            else:
                print(f"   Note: python-docx2pdf only works on Windows (current OS: {current_os})")
            # Continue to fallback methods

        except Exception as e:
            # Conversion failed - log error and try fallback
            print(f"   ⚠️  python-docx2pdf conversion failed: {e}")
            if "com_error" in str(type(e).__name__).lower():
                print("      Hint: Is Microsoft Word installed and licensed?")
            # Continue to fallback methods

        # Strategy 2: Try LibreOffice command-line conversion (cross-platform)
        # This is the primary method for Linux/macOS and fallback for Windows
        # LibreOffice provides excellent compatibility across platforms
        print("\n   Trying fallback: LibreOffice (cross-platform)")
        print("   Quality: Very Good (~95% formatting preserved)")

        try:
            import subprocess

            # Detect LibreOffice installation
            # Different platforms use different executable names
            libreoffice_commands = []

            if current_os == 'Windows':
                # Windows: Check common installation paths
                libreoffice_commands = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    "soffice.exe",  # If in PATH
                ]
            elif current_os == 'Darwin':  # macOS
                libreoffice_commands = [
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                    "soffice",  # If in PATH
                ]
            else:  # Linux and others
                libreoffice_commands = [
                    "libreoffice",  # Standard command
                    "soffice",      # Alternative command
                    "/usr/bin/libreoffice",
                    "/usr/bin/soffice",
                ]

            # Try each possible LibreOffice command
            libreoffice_cmd = None
            for cmd in libreoffice_commands:
                try:
                    # Test if command exists and is executable
                    # Use 'where' on Windows, 'which' on Unix
                    if current_os == 'Windows':
                        result = subprocess.run(
                            ['where', cmd if not cmd.startswith('C:') else cmd.split('\\')[-1]],
                            capture_output=True,
                            timeout=5
                        )
                    else:
                        result = subprocess.run(
                            ['which', cmd.split('/')[-1]],
                            capture_output=True,
                            timeout=5
                        )

                    if result.returncode == 0 or os.path.exists(cmd):
                        libreoffice_cmd = cmd
                        break
                except Exception:
                    # Command not found, try next
                    continue

            if libreoffice_cmd is None:
                # LibreOffice not found - provide installation instructions
                print("   ❌ LibreOffice not found on system")
                print("\n   📥 Install LibreOffice for PDF export:")
                if current_os == 'Windows':
                    print("      Download: https://www.libreoffice.org/download/download/")
                    print("      Or via Chocolatey: choco install libreoffice")
                elif current_os == 'Darwin':
                    print("      brew install --cask libreoffice")
                elif current_os == 'Linux':
                    print("      Ubuntu/Debian: sudo apt install libreoffice-writer")
                    print("      Fedora/RHEL:   sudo dnf install libreoffice-writer")
                    print("      Arch:          sudo pacman -S libreoffice-fresh")
                else:
                    print("      Visit: https://www.libreoffice.org/download/download/")

                # No more fallbacks available
                print("\n   ℹ️  Alternatively, open the .docx file and use File → Export as PDF")
                return None

            print(f"   Found LibreOffice: {libreoffice_cmd}")

            # Construct LibreOffice conversion command
            # --headless: Run without GUI
            # --convert-to pdf: Output format
            # --outdir: Specify output directory
            output_dir = pdf_path.parent

            cmd_args = [
                libreoffice_cmd,
                '--headless',  # No GUI
                '--convert-to', 'pdf',  # Target format
                '--outdir', str(output_dir),  # Output directory
                docx_path  # Input file
            ]

            print(f"   Running: {' '.join(cmd_args[:4])} ...")

            # Execute conversion with timeout to prevent hanging
            # LibreOffice conversion typically takes 5-30 seconds depending on doc size
            result = subprocess.run(
                cmd_args,
                capture_output=True,
                text=True,
                timeout=120  # 2 minute timeout (generous for large documents)
            )

            # Check if conversion succeeded
            if result.returncode == 0 and pdf_path.exists():
                # Verify PDF file is not empty
                pdf_size = pdf_path.stat().st_size
                if pdf_size > 0:
                    print(f"✅ PDF successfully created via LibreOffice: {pdf_path.absolute()}")
                    print(f"   Size: {pdf_size:,} bytes")
                    return str(pdf_path.absolute())
                else:
                    print("⚠️  LibreOffice created empty PDF file")
            else:
                # Conversion failed - print diagnostic info
                print(f"   ❌ LibreOffice conversion failed (exit code: {result.returncode})")
                if result.stderr:
                    print(f"   Error output: {result.stderr[:200]}")
                if result.stdout:
                    print(f"   Standard output: {result.stdout[:200]}")

        except subprocess.TimeoutExpired:
            print("   ❌ LibreOffice conversion timed out (>2 minutes)")
            print("      Hint: Document may be too large or complex")
        except Exception as e:
            print(f"   ❌ LibreOffice conversion error: {e}")

        # Strategy 3: reportlab fallback (future enhancement)
        # Note: Not implemented in this version due to complexity of converting
        # python-docx Document object to reportlab canvas while preserving formatting
        print("\n   ℹ️  All automatic PDF conversion methods failed")
        print("   📝 Manual conversion options:")
        print("      1. Open documentation.docx in Word/LibreOffice")
        print("      2. Use File → Export as PDF")
        print("      3. Or try online converter: https://www.ilovepdf.com/word_to_pdf")

        return None

class DocumentationGenerator:
    """Main orchestrator"""
    
    def __init__(self):
        self.project_name = os.getenv('PROJECT_NAME', 'My Project')
        self.project_path = Path(os.getenv('PROJECT_PATH', '.'))

        self.gemini_agent = GeminiDocAgent()
        self.screenshot_agent = ScreenshotAgent()
        self.mermaid_agent = MermaidAgent()
        self.assembler = DocumentAssembler()

        # Optimization settings
        self.max_sections = int(os.getenv('MAX_SECTIONS', '15'))
        self.min_sections = int(os.getenv('MIN_SECTIONS', '9'))
        self.max_screenshots = int(os.getenv('MAX_SCREENSHOTS_PER_DOCUMENT', '8'))
        self.max_mermaid_diagrams = int(os.getenv('MAX_MERMAID_DIAGRAMS', '3'))
        self.screenshot_priority = os.getenv('SCREENSHOT_PRIORITY_SECTIONS',
                                            'installation,configuration,architecture,usage').lower().split(',')
        self.screenshot_priority = [s.strip() for s in self.screenshot_priority]

        # Tracking counters
        self.screenshot_count = 0
        self.mermaid_count = 0
    
    def load_context(self) -> str:
        """Load project context from repomix or scan directory"""
        use_repomix = os.getenv('USE_REPOMIX', 'true').lower() == 'true'
        repomix_file = os.getenv('REPOMIX_FILE_PATH')
        
        if use_repomix and repomix_file and os.path.exists(repomix_file):
            print(f"📖 Loading context from repomix file: {repomix_file}")
            with open(repomix_file, 'r', encoding='utf-8') as f:
                context = f.read()
            print(f"✓ Loaded {len(context)} characters from repomix")
            return context
        
        # Fallback: scan directory
        print("📖 Scanning project directory...")
        context = f"Project: {self.project_name}\n\n"
        
        # Read README
        for readme in ['README.md', 'readme.md', 'README.txt', 'README']:
            readme_path = self.project_path / readme
            if readme_path.exists():
                try:
                    with open(readme_path, 'r', encoding='utf-8') as f:
                        context += f"=== README ===\n{f.read()}\n\n"
                    print(f"✓ Found {readme}")
                    break
                except:
                    pass
        
        # Scan directory structure
        excluded = os.getenv('EXCLUDED_DIRECTORIES', '').split(',')
        excluded = [d.strip() for d in excluded if d.strip()]
        
        context += "=== Project Structure ===\n"
        file_count = 0
        
        for item in self.project_path.rglob('*'):
            if any(ex in str(item) for ex in excluded):
                continue
            
            if item.is_file():
                rel_path = item.relative_to(self.project_path)
                context += f"- {rel_path}\n"
                
                # Read small code files
                max_size = int(os.getenv('MAX_FILE_SIZE_KB', '100')) * 1024
                if item.stat().st_size < max_size and item.suffix in ['.py', '.js', '.java', '.cpp', '.go', '.rs']:
                    try:
                        with open(item, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            context += f"\n=== File: {rel_path} ===\n{file_content}\n\n"
                            file_count += 1
                            
                            if file_count >= 20:  # Limit files
                                break
                    except:
                        pass
        
        print(f"✓ Scanned project, loaded {file_count} files")
        return context

    def _should_capture_screenshot(self, section: DocumentSection) -> bool:
        """Determine if a section should have screenshots based on priority.

        Args:
            section: The DocumentSection to evaluate

        Returns:
            bool: True if screenshots should be captured for this section
        """
        section_title_lower = section.title.lower()

        # Check if any priority keyword is in the section title
        for priority_keyword in self.screenshot_priority:
            if priority_keyword in section_title_lower:
                return True

        return False

    def generate(self):
        """Main generation pipeline"""
        print(f"\n{'='*60}")
        print(f"🚀 Starting documentation generation for: {self.project_name}")
        print(f"{'='*60}\n")
        
        # Phase 1: Load context
        print("📖 Phase 1: Loading project context...")
        context = self.load_context()
        print(f"✓ Loaded {len(context)} characters of context\n")
        
        # Phase 2: Create plan
        print("📋 Phase 2: Creating documentation plan...")
        print(f"   Configuration: {self.min_sections}-{self.max_sections} sections based on project size")
        plan = self.gemini_agent.create_documentation_plan(
            context, self.project_name,
            min_sections=self.min_sections,
            max_sections=self.max_sections
        )
        print(f"✓ Created plan with {len(plan.sections)} sections:")
        for i, section in enumerate(plan.sections, 1):
            print(f"  {i}. {section.title}")
        print()
        
        # Phase 3: Generate content
        print("✍️  Phase 3: Generating content...")
        print(f"   Screenshot limit: {self.max_screenshots} per document")
        previous_content = ""
        enable_screenshots = os.getenv('ENABLE_SCREENSHOTS', 'true').lower() == 'true'

        for i, section in enumerate(plan.sections, 1):
            print(f"  [{i}/{len(plan.sections)}] Generating: {section.title}")

            content = self.gemini_agent.generate_section_content(
                section, context, previous_content
            )
            section.content = content
            previous_content += f"\n\n## {section.title}\n{content}"

            # Capture screenshots if enabled - with optimization
            if enable_screenshots and section.images and self.screenshot_count < self.max_screenshots:
                # Check if this section is in priority list
                should_capture = self._should_capture_screenshot(section)

                if should_capture:
                    remaining_screenshots = self.max_screenshots - self.screenshot_count
                    screenshots_to_capture = min(len(section.images), remaining_screenshots)

                    if screenshots_to_capture > 0:
                        print(f"      📸 Capturing {screenshots_to_capture} screenshot(s) ({self.screenshot_count}/{self.max_screenshots} used)...")
                        targets = self.gemini_agent.identify_screenshot_targets(section, context)

                        for j, target in enumerate(targets):
                            if j >= screenshots_to_capture:
                                break

                            path = None
                            if target['target_type'] == 'code_file':
                                path = self.screenshot_agent.capture_code_file(
                                    target['target_path'],
                                    target.get('instructions', '')
                                )
                            elif target['target_type'] == 'directory_structure':
                                path = self.screenshot_agent.capture_directory_tree()

                            if path:
                                section.images[j]['path'] = path
                                self.screenshot_count += 1
                else:
                    print(f"      ⏭️  Skipping screenshots (not in priority sections)")
            elif enable_screenshots and section.images and self.screenshot_count >= self.max_screenshots:
                print(f"      ⚠️  Screenshot limit reached ({self.max_screenshots}), skipping")

        print(f"✓ Content generation complete ({self.screenshot_count} screenshots captured)\n")

        # Phase 3.4: Generate Mermaid diagrams for architecture sections
        enable_mermaid = os.getenv('ENABLE_MERMAID_DIAGRAMS', 'true').lower() == 'true'
        if enable_mermaid and self.mermaid_agent.use_mermaid:
            print("🎨 Phase 3.4: Generating architecture diagrams...")
            print(f"   Diagram limit: {self.max_mermaid_diagrams} for entire document")

            # Priority sections for Mermaid diagrams (in order of priority)
            diagram_priority_keywords = ['overview', 'architecture', 'design', 'components']

            # Sort sections by priority
            diagram_candidates = []
            for section in plan.sections:
                section_title_lower = section.title.lower()
                for priority_idx, keyword in enumerate(diagram_priority_keywords):
                    if keyword in section_title_lower:
                        diagram_candidates.append((priority_idx, section))
                        break

            # Sort by priority (lower priority_idx = higher priority)
            diagram_candidates.sort(key=lambda x: x[0])

            # Generate diagrams up to the limit
            for priority_idx, section in diagram_candidates:
                if self.mermaid_count >= self.max_mermaid_diagrams:
                    print(f"    ⚠️  Diagram limit reached ({self.max_mermaid_diagrams}), skipping remaining sections")
                    break

                print(f"    Generating diagram for: {section.title} ({self.mermaid_count + 1}/{self.max_mermaid_diagrams})")

                # Determine diagram type based on section
                diagram_type = "flowchart"
                if "component" in section.title.lower():
                    diagram_type = "flowchart"
                elif "architecture" in section.title.lower():
                    diagram_type = "flowchart"

                # Generate Mermaid code
                description = f"System architecture diagram for {section.title}"
                mermaid_code = self.mermaid_agent.generate_diagram_code(
                    self.gemini_agent, context, diagram_type, description
                )

                if mermaid_code:
                    # Render to PNG
                    diagram_name = section.title.lower().replace(' ', '_')
                    diagram_path = self.mermaid_agent.render_diagram(mermaid_code, diagram_name)

                    if diagram_path:
                        # Add diagram to section images
                        if not section.images:
                            section.images = []
                        section.images.insert(0, {
                            'description': f'Architecture Diagram: {section.title}',
                            'path': diagram_path
                        })
                        # Store mermaid diagram info
                        section.mermaid_diagrams.append({
                            'description': description,
                            'code': mermaid_code,
                            'path': diagram_path
                        })
                        self.mermaid_count += 1

            print(f"✓ Architecture diagrams generated ({self.mermaid_count} total)\n")

        # Phase 3.5: Capture live app screenshots
        live_app_enabled = os.getenv('LIVE_APP_ENABLED', 'false').lower() == 'true'
        if live_app_enabled:
            print("📸 Phase 3.5: Capturing live application screenshots...")
            
            # Collect all LIVE_APP_URL_* environment variables
            live_urls = {}
            for key, value in os.environ.items():
                if key.startswith('LIVE_APP_URL_'):
                    name = key.replace('LIVE_APP_URL_', '').lower()
                    live_urls[name] = value
            
            if live_urls:
                print(f"  Found {len(live_urls)} URLs to capture")
                for name, url in live_urls.items():
                    print(f"    Capturing: {name} -> {url}")
                    screenshot_path = self.screenshot_agent.capture_live_url(url, name)
                    
                    if screenshot_path:
                        # Add to first relevant section
                        for section in plan.sections:
                            if any(keyword in section.title.lower() 
                                   for keyword in ['overview', 'interface', 'usage', 'introduction']):
                                if not section.images:
                                    section.images = []
                                section.images.append({
                                    'description': f'Application Screenshot: {name}',
                                    'path': screenshot_path
                                })
                                break
                print("✓ Live app screenshots captured\n")
            else:
                print("  No live app URLs configured\n")
        
        # Phase 4: Assemble document
        print("📄 Phase 4: Assembling Word document...")
        self.assembler.add_title_page(plan.title, self.project_name)
        
        for section in plan.sections:
            self.assembler.add_section(section)
        
        output_path = self.assembler.save()

        # Phase 5: Export to PDF (if enabled)
        # Why separate phase: PDF export is optional and may fail without affecting
        # main .docx generation. Keeps concerns separated.
        pdf_path = self.assembler.save_as_pdf()

        print(f"\n{'='*60}")
        print("✅ Documentation generation complete!")
        print(f"{'='*60}")
        print(f"\n📄 Document: {output_path}")

        # Display PDF path if generated successfully
        if pdf_path:
            print(f"📄 PDF Export: {pdf_path}")

        if enable_screenshots:
            screenshot_dir = os.getenv('SCREENSHOTS_DIRECTORY', './screenshots')
            print(f"📸 Screenshots: {Path(screenshot_dir).absolute()}")

        print("\n💡 Next steps:")
        print("  1. Open the .docx file in Word/Google Docs/LibreOffice")
        print("  2. Review and edit as needed")
        print("  3. Share with your team!\n")

        return output_path


def main():
    """Entry point"""
    try:
        generator = DocumentationGenerator()
        generator.generate()
    except KeyboardInterrupt:
        print("\n\n⚠️  Generation interrupted by user")
        return 1
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())